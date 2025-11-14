# SISTEMA INTELIGENTE DE ANÁLISIS DE CASOS DE PRUEBA
# ¡CON HISTORIAL EN BASE DE DATOS!

from flask import render_template, flash, redirect, url_for, request, current_app, session, send_file
from flask_login import current_user, login_required
from app import db
from app.analysis import bp
from app.analysis.forms import RequerimientoUploadForm
# ¡Importamos el nuevo modelo Analisis!
from app.models import Plantilla, MapaPlantilla, Analisis
import os
import docx
from docx.shared import Pt 
import google.generativeai as genai
import json
import re
from io import BytesIO
import openpyxl
import traceback
# Ya no necesitamos tempfile ni pickle

# --- Constantes ---
ALLOWED_EXTENSIONS = {'xlsx', 'docx'}
TAG_REGEX = re.compile(r"(\{\{.*?\}\})") 
TAG_CLEANER = re.compile(r"\{\{(.*?)\}\}") 

# --- ¡FUNCIONES tempfile ELIMINADAS! ---

# ==============================================================================
# FUNCIÓN DE ANÁLISIS INTELIGENTE (Sin cambios)
# ==============================================================================
def analizar_complejidad_requerimiento(texto_requerimiento):
    """
    Analiza el requerimiento y determina automáticamente la cantidad óptima
    de casos de prueba necesarios basándose en múltiples factores.
    """
    
    texto_lower = texto_requerimiento.lower()
    
    # 1. MÉTRICAS BÁSICAS
    palabras = texto_requerimiento.split()
    total_palabras = len(palabras)
    lineas = texto_requerimiento.split('\n')
    total_lineas = len([l for l in lineas if l.strip()])
    
    # 2. DETECTAR CRITERIOS DE ACEPTACIÓN (Lógica mejorada)
    patron_ca = re.compile(r'^(ca|ac|criterio)[\s\-:]*\d*.*$', re.IGNORECASE | re.MULTILINE)
    patron_historia = re.compile(r'como\s+.*quiero\s+.*para\s+.*', re.IGNORECASE | re.DOTALL)
    patron_gherkin = re.compile(r'^(escenario|dado|given|cuando|when).*$', re.IGNORECASE | re.MULTILINE)

    criterios_encontrados = 0
    criterios_encontrados += len(re.findall(patron_ca, texto_requerimiento))
    criterios_encontrados += len(re.findall(patron_historia, texto_requerimiento))
    criterios_encontrados += len(re.findall(patron_gherkin, texto_requerimiento))

    if criterios_encontrados == 0:
        listas_numeradas = len(re.findall(r'^\s*\d+[\.\)]\s+', texto_requerimiento, re.MULTILINE))
        listas_viñetas = len(re.findall(r'^\s*[-•*]\s+', texto_requerimiento, re.MULTILINE))
        criterios_encontrados = listas_numeradas + listas_viñetas

    if criterios_encontrados == 0 and total_palabras > 50:
        criterios_encontrados = 1
    
    # 3. DETECTAR PALABRAS CLAVE DE COMPLEJIDAD
    keywords_validacion = ['validar', 'verificar', 'validación', 'verificación', 'comprobar', 'asegurar', 'garantizar']
    keywords_condicionales = ['si', 'cuando', 'entonces', 'caso contrario', 'de lo contrario', 'if', 'when', 'then', 'else', 'otherwise']
    keywords_campos = ['campo', 'formulario', 'input', 'entrada', 'dato', 'textbox', 'checkbox', 'dropdown', 'select', 'botón', 'button']
    keywords_estados = ['estado', 'status', 'activo', 'inactivo', 'pendiente', 'aprobado', 'rechazado', 'completado']
    keywords_errores = ['error', 'excepción', 'fallo', 'incorrecto', 'inválido', 'mensaje de error', 'alerta', 'warning']
    keywords_flujo = ['flujo', 'proceso', 'paso', 'secuencia', 'etapa', 'workflow', 'navigation']
    keywords_integracion = ['integración', 'api', 'servicio', 'base de datos', 'endpoint', 'request', 'response', 'conexión']
    keywords_seguridad = ['seguridad', 'autenticación', 'autorización', 'permisos', 'roles', 'token', 'sesión', 'login', 'logout', 'contraseña']
    
    count_validacion = sum(1 for k in keywords_validacion if k in texto_lower)
    count_condicionales = sum(1 for k in keywords_condicionales if k in texto_lower)
    count_campos = sum(1 for k in keywords_campos if k in texto_lower)
    count_estados = sum(1 for k in keywords_estados if k in texto_lower)
    count_errores = sum(1 for k in keywords_errores if k in texto_lower)
    count_flujo = sum(1 for k in keywords_flujo if k in texto_lower)
    count_integracion = sum(1 for k in keywords_integracion if k in texto_lower)
    count_seguridad = sum(1 for k in keywords_seguridad if k in texto_lower)
    
    # 4. CALCULAR PUNTUACIÓN DE COMPLEJIDAD
    complejidad_score = 0
    
    if total_palabras < 100:
        complejidad_score += 5
    elif total_palabras < 300:
        complejidad_score += 10
    elif total_palabras < 600:
        complejidad_score += 15
    elif total_palabras < 1000:
        complejidad_score += 20
    else:
        complejidad_score += 25
    
    complejidad_score += min(criterios_encontrados * 5, 30)
    complejidad_score += min(count_validacion * 2, 10)
    complejidad_score += min(count_condicionales * 2, 10)
    complejidad_score += min(count_campos * 1.5, 8)
    complejidad_score += min(count_estados * 1.5, 8)
    complejidad_score += min(count_errores * 2, 12)
    complejidad_score += min(count_flujo * 1.5, 8)
    complejidad_score += min(count_integracion * 2.5, 15)
    complejidad_score += min(count_seguridad * 3, 18)
    
    # 5. DETERMINAR CANTIDAD DE CASOS SEGÚN SCORE
    if complejidad_score < 20:
        cantidad_casos = 5
        nivel = "Muy Simple"
    elif complejidad_score < 35:
        cantidad_casos = 8
        nivel = "Simple"
    elif complejidad_score < 50:
        cantidad_casos = 12
        nivel = "Estándar"
    elif complejidad_score < 70:
        cantidad_casos = 18
        nivel = "Complejo"
    elif complejidad_score < 90:
        cantidad_casos = 25
        nivel = "Muy Complejo"
    else:
        cantidad_casos = 35
        nivel = "Extremadamente Complejo"
    
    # 6. AJUSTE INTELIGENTE
    if criterios_encontrados > 5:
        cantidad_casos = max(cantidad_casos, criterios_encontrados * 2)
    
    if count_integracion > 2 or count_seguridad > 2:
        cantidad_casos = int(cantidad_casos * 1.3)
    
    if count_estados > 3 or count_flujo > 3:
        cantidad_casos = int(cantidad_casos * 1.2)
    
    cantidad_casos = min(cantidad_casos, 50)
    cantidad_casos = max(cantidad_casos, 5)
    
    # 7. PREPARAR RESULTADO
    resultado = {
        'cantidad_casos': cantidad_casos,
        'nivel_complejidad': nivel,
        'score': round(complejidad_score, 2),
        'metricas': {
            'palabras': total_palabras,
            'lineas': total_lineas,
            'criterios_aceptacion': criterios_encontrados,
            'listas': 0
        },
        'factores': {
            'validaciones': count_validacion,
            'condicionales': count_condicionales,
            'campos_formulario': count_campos,
            'estados': count_estados,
            'manejo_errores': count_errores,
            'flujos': count_flujo,
            'integracion': count_integracion,
            'seguridad': count_seguridad
        },
        'recomendacion': f"Se generarán {cantidad_casos} casos de prueba para una cobertura óptima."
    }
    
    return resultado

# ==============================================================================
# FUNCIONES AUXILIARES (Sin cambios)
# ==============================================================================
def leer_texto_requerimiento(archivo):
    # ... (Tu código existente)
    filename = archivo.filename
    text_content = ""
    try:
        if filename.endswith('.docx'):
            doc = docx.Document(archivo)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        elif filename.endswith('.txt') or filename.endswith('.md'):
            text_content = archivo.read().decode('utf-8')
        return text_content
    except Exception as e:
        flash(f"Error al leer el archivo: {str(e)}", "danger")
        return None

def generar_excel_entregable(plantilla_obj, datos_ia):
    # ... (Tu código existente)
    try:
        path_plantilla = os.path.join(current_app.config['UPLOAD_FOLDER'], plantilla_obj.filename_seguro)
        workbook = openpyxl.load_workbook(path_plantilla)
        if not plantilla_obj.sheet_name or not plantilla_obj.header_row:
            print(f"Error: Plantilla '{plantilla_obj.nombre_plantilla}' (ID: {plantilla_obj.id}) no está mapeada.")
            return None
        sheet = workbook[plantilla_obj.sheet_name]
        fila_inicio = plantilla_obj.header_row + 1
        mapa_etiquetas = {mapa.etiqueta: mapa.coordenada for mapa in plantilla_obj.mapas}
        if all(m.tipo_mapa == 'fila_tabla' for m in plantilla_obj.mapas):
            for caso in datos_ia:
                for etiqueta, valor in caso.items():
                    if etiqueta in mapa_etiquetas:
                        columna = mapa_etiquetas[etiqueta]
                        valor_final = valor
                        if isinstance(valor, list):
                            valor_final = "\n".join(str(v) for v in valor)
                        sheet[f"{columna}{fila_inicio}"] = valor_final
                fila_inicio += 1
        else:
            if datos_ia:
                caso_principal = datos_ia[0]
                for etiqueta, valor in caso_principal.items():
                    if etiqueta in mapa_etiquetas:
                        coordenada = mapa_etiquetas[etiqueta]
                        valor_final = valor
                        if isinstance(valor, list):
                            valor_final = "\n".join(str(v) for v in valor)
                        sheet[coordenada] = valor_final
        buffer_memoria = BytesIO()
        workbook.save(buffer_memoria)
        buffer_memoria.seek(0)
        return buffer_memoria
    except Exception as e:
        print(f"Error generando Excel: {e}")
        traceback.print_exc()
        return None

def reemplazar_texto_en_parrafo(parrafo, etiqueta, valor):
    # ... (Tu código existente)
    if etiqueta in parrafo.text:
        for run in parrafo.runs:
            if etiqueta in run.text:
                run.text = run.text.replace(etiqueta, str(valor))

def generar_word_entregable(plantilla_obj, datos_ia):
    # ... (Tu código existente)
    try:
        path_plantilla = os.path.join(current_app.config['UPLOAD_FOLDER'], plantilla_obj.filename_seguro)
        document = docx.Document(path_plantilla)
        mapa_simple = {mapa.etiqueta: f"{{{{{mapa.etiqueta}}}}}" for mapa in plantilla_obj.mapas if mapa.tipo_mapa == 'celda_simple'}
        mapa_tabla = {mapa.etiqueta: int(mapa.coordenada) for mapa in plantilla_obj.mapas if mapa.tipo_mapa == 'fila_tabla'}
        etiquetas_tabla_keys = list(mapa_tabla.keys())
        for i, caso_ia in enumerate(datos_ia):
            if i == 0:
                for etiqueta_limpia, etiqueta_sucia in mapa_simple.items():
                    if etiqueta_limpia in caso_ia:
                        valor = caso_ia[etiqueta_limpia]
                        for para in document.paragraphs: reemplazar_texto_en_parrafo(para, etiqueta_sucia, valor)
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs: reemplazar_texto_en_parrafo(para, etiqueta_sucia, valor)
                if etiquetas_tabla_keys:
                    tabla_encontrada = None
                    fila_plantilla_idx = -1
                    for table in document.tables:
                        if fila_plantilla_idx != -1: break
                        for r_idx, row in enumerate(table.rows):
                            try:
                                primera_etiqueta = etiquetas_tabla_keys[0]
                                primera_celda_idx = mapa_tabla[primera_etiqueta]
                                if f"{{{{{primera_etiqueta}}}}}" in row.cells[primera_celda_idx].text:
                                    tabla_encontrada = table
                                    fila_plantilla_idx = r_idx
                                    break
                            except (IndexError, KeyError): continue
                    if tabla_encontrada and fila_plantilla_idx != -1:
                        datos_pasos = caso_ia.get("PASOS", [])
                        if not datos_pasos and all(key in caso_ia for key in etiquetas_tabla_keys): datos_pasos = [caso_ia]
                        if datos_pasos and len(datos_pasos) > 0:
                            fila_plantilla_original = tabla_encontrada.rows[fila_plantilla_idx]
                            textos_plantilla = {}
                            for etiqueta_limpia, col_idx in mapa_tabla.items():
                                try: textos_plantilla[etiqueta_limpia] = fila_plantilla_original.cells[col_idx].text
                                except IndexError: textos_plantilla[etiqueta_limpia] = f"{{{{{etiqueta_limpia}}}}}"
                            for paso_idx, paso_ia in enumerate(datos_pasos):
                                if paso_idx == 0: fila_actual = fila_plantilla_original
                                else: fila_actual = tabla_encontrada.add_row()
                                for etiqueta_limpia, col_idx in mapa_tabla.items():
                                    try:
                                        if paso_idx > 0: fila_actual.cells[col_idx].text = textos_plantilla.get(etiqueta_limpia, "")
                                        if etiqueta_limpia in paso_ia:
                                            valor_paso = str(paso_ia[etiqueta_limpia])
                                            etiqueta_sucia = f"{{{{{etiqueta_limpia}}}}}"
                                            texto_actual = fila_actual.cells[col_idx].text
                                            fila_actual.cells[col_idx].text = texto_actual.replace(etiqueta_sucia, valor_paso)
                                    except IndexError: continue
            else:
                document.add_page_break()
                id_caso_val = caso_ia.get('ID del caso de prueba', caso_ia.get('ID_CASO_PRUEBA', ''))
                titulo_val = caso_ia.get('Descripción del caso de prueba', caso_ia.get('TITULO_CASO_PRUEBA', ''))
                titulo_completo = f"Caso de Prueba: {id_caso_val} - {titulo_val}"
                document.add_heading(titulo_completo, level=2)
                for etiqueta_limpia in mapa_simple.keys():
                    if etiqueta_limpia in caso_ia:
                        document.add_heading(etiqueta_limpia.replace('_', ' ').title(), level=3)
                        document.add_paragraph(str(caso_ia[etiqueta_limpia]))
                datos_pasos = caso_ia.get("PASOS", [])
                if not datos_pasos and all(key in caso_ia for key in etiquetas_tabla_keys): datos_pasos = [caso_ia]
                if datos_pasos and etiquetas_tabla_keys:
                    document.add_heading("Detalle de Ejecución", level=3)
                    tabla_nueva = document.add_table(rows=1, cols=len(etiquetas_tabla_keys))
                    try: tabla_nueva.style = 'Light Grid Accent 1'
                    except KeyError: tabla_nueva.style = 'Table Grid'
                    hdr_cells = tabla_nueva.rows[0].cells
                    for idx, etiqueta in enumerate(etiquetas_tabla_keys): hdr_cells[idx].text = etiqueta
                    for paso_ia in datos_pasos:
                        row_cells = tabla_nueva.add_row().cells
                        for idx, etiqueta_limpia in enumerate(etiquetas_tabla_keys):
                            if etiqueta_limpia in paso_ia: row_cells[idx].text = str(paso_ia[etiqueta_limpia])
                            else: row_cells[idx].text = ""
        buffer_memoria = BytesIO()
        document.save(buffer_memoria)
        buffer_memoria.seek(0)
        return buffer_memoria
    except Exception as e:
        print(f"Error generando Word: {e}")
        traceback.print_exc()
        return None

# ==============================================================================
# RUTAS REFACTORIZADAS (Versión "Historial en BD" + Fix de IA)
# ==============================================================================

@bp.route('/analysis', methods=['GET', 'POST'])
@login_required
def analysis_index():
    """
    Ruta principal para análisis.
    Maneja tanto el envío (POST) como la visualización (GET).
    """
    form = RequerimientoUploadForm()
    
    plantillas_usuario = Plantilla.query.filter_by(
        autor=current_user
    ).with_entities(Plantilla.id, Plantilla.nombre_plantilla).all()
    form.plantilla.choices = [(p.id, p.nombre_plantilla) for p in plantillas_usuario]
    
    view_id = request.args.get('view_id', None)
    
    analisis_obj = None
    analisis_info = None
    ai_result_data = None
    texto_requerimiento_raw = None
    ai_result_raw_text = None 

    if view_id:
        analisis_obj = Analisis.query.get_or_404(view_id)
        if analisis_obj.autor != current_user:
            flash("Acceso no autorizado.", "danger")
            return redirect(url_for('analysis.analysis_index'))
            
        analisis_info = {
            'nivel': analisis_obj.nivel_complejidad,
            'casos': analisis_obj.casos_generados,
            'criterios': analisis_obj.criterios_detectados,
            'palabras': analisis_obj.palabras_analizadas
        }
        texto_requerimiento_raw = analisis_obj.texto_requerimiento_raw
        ai_result_raw_text = analisis_obj.ai_result_json
        
        if analisis_obj.ai_result_json:
            try:
                match = re.search(r'\[.*\]', analisis_obj.ai_result_json, re.DOTALL)
                if match:
                    ai_result_data = json.loads(match.group(0))
            except Exception as e:
                print(f"Error parseando JSON de BD: {e}")
                ai_result_data = None
    
    if form.validate_on_submit():
        archivo = form.archivo_requerimiento.data
        plantilla_seleccionada_id = form.plantilla.data
        plantilla_obj = Plantilla.query.get(plantilla_seleccionada_id)
        
        mapas = plantilla_obj.mapas.all()
        if not mapas:
            flash(f"La plantilla '{plantilla_obj.nombre_plantilla}' no tiene etiquetas escaneadas.", "danger")
            return redirect(url_for('analysis.analysis_index'))
        if plantilla_obj.tipo_archivo == 'Excel' and (not plantilla_obj.sheet_name or not plantilla_obj.header_row):
             flash(f"La plantilla de Excel '{plantilla_obj.nombre_plantilla}' está incompleta.", "danger")
             return redirect(url_for('core.map_step_1_sheet', plantilla_id=plantilla_obj.id))
            
        etiquetas_simples = [m.etiqueta for m in mapas if m.tipo_mapa == 'celda_simple']
        etiquetas_tabla = [m.etiqueta for m in mapas if m.tipo_mapa == 'fila_tabla']
        
        texto_requerimiento = leer_texto_requerimiento(archivo)
        
        if texto_requerimiento:
            print("\n" + "="*80)
            print("🧠 ANALIZANDO COMPLEJIDAD DEL REQUERIMIENTO...")
            print("="*80)
            analisis = analizar_complejidad_requerimiento(texto_requerimiento)
            num_casos = analisis['cantidad_casos']
            # (impresión en consola de métricas)
            
            try:
                api_key = os.environ.get('GOOGLE_API_KEY') 
                if not api_key:
                    flash("Error: GOOGLE_API_KEY no está configurada.", "danger")
                    return redirect(url_for('analysis.analysis_index'))
                
                genai.configure(api_key=api_key)
                
                # Volvemos al modelo FLASH que te funciona
                model = genai.GenerativeModel('gemini-2.0-flash-exp') 
                
                # Prompt SIMPLIFICADO para el modelo FLASH
                prompt = f"""
Eres un asistente de QA que genera casos de prueba en formato JSON.
Basado en el siguiente requerimiento, genera EXACTAMENTE {num_casos} casos de prueba.

REQUERIMIENTO A ANALIZAR:
---
{texto_requerimiento}
---

INSTRUCCIONES CRÍTICAS:
1. La respuesta debe ser ÚNICAMENTE un array JSON válido, sin texto adicional.
2. Debes generar EXACTAMENTE {num_casos} casos de prueba.
3. Basa tus casos en los Criterios de Aceptación y el texto del requerimiento.
"""
                
                if etiquetas_simples and etiquetas_tabla:
                    prompt += f"""
4. Cada caso debe tener estos campos principales:
   {json.dumps(etiquetas_simples, ensure_ascii=False)}
   
5. ADICIONALMENTE, cada caso debe tener "PASOS" (array de objetos):
   {json.dumps(etiquetas_tabla, ensure_ascii=False)}
   
6. Campos principales = strings. "PASOS" = array con 2-10 pasos.
"""
                elif not etiquetas_simples and etiquetas_tabla:
                    prompt += f"""
4. Cada caso debe tener estos campos:
   {json.dumps(etiquetas_tabla, ensure_ascii=False)}

5. Todos strings. Si hay campo "Pasos", usar saltos de línea numerados.
"""
                elif etiquetas_simples and not etiquetas_tabla:
                    prompt += f"""
4. Cada caso debe tener estos campos:
   {json.dumps(etiquetas_simples, ensure_ascii=False)}

5. Todos strings. Si hay campo "Pasos", usar saltos de línea numerados.
"""

                prompt += f"""

IMPORTANTE: 
- Genera TODOS los campos exactamente como están escritos
- NO inventes campos adicionales
- Asegura que el JSON sea válido y parseable
- EXACTAMENTE {num_casos} casos, ni más ni menos
Responde ÚNICAMENTE con el array JSON: [ ... ]
"""
                
                print("🤖 Enviando prompt a Gemini...")
                response = model.generate_content(prompt)
                
                ai_json_string = None
                casos_generados_count = 0
                try:
                    match = re.search(r'\[.*\]', response.text, re.DOTALL)
                    if match:
                        ai_json_string = match.group(0) 
                        casos_generados = json.loads(ai_json_string)
                        casos_generados_count = len(casos_generados)
                        print(f"✅ Generación exitosa: {casos_generados_count} casos")
                        flash(f"✅ Análisis completado: {casos_generados_count} casos generados", "success")
                    else:
                        flash("⚠️ Respuesta no parseable. Intenta nuevamente.", "warning")
                        ai_json_string = response.text 
                except json.JSONDecodeError as je:
                    print(f"❌ Error JSON: {je}")
                    flash("Error: Respuesta de IA inválida.", "danger")
                    ai_json_string = response.text 

                nuevo_analisis = Analisis(
                    autor=current_user,
                    plantilla_usada=plantilla_obj,
                    nombre_requerimiento=archivo.filename,
                    texto_requerimiento_raw=texto_requerimiento,
                    nivel_complejidad=analisis['nivel_complejidad'],
                    casos_generados=casos_generados_count,
                    criterios_detectados=analisis['metricas']['criterios_aceptacion'],
                    palabras_analizadas=analisis['metricas']['palabras'],
                    ai_result_json=ai_json_string 
                )
                
                db.session.add(nuevo_analisis)
                db.session.commit()
                
                print(f"💾 Análisis {nuevo_analisis.id} guardado en la base de datos.")
                
                return redirect(url_for('analysis.analysis_index', view_id=nuevo_analisis.id))
                
            except Exception as e:
                flash(f"Error con API Gemini: {str(e)}", "danger")
                traceback.print_exc()
        
        return redirect(url_for('analysis.analysis_index')) 

    historial_analisis = Analisis.query.filter_by(
        autor=current_user
    ).order_by(Analisis.timestamp.desc()).all()

    return render_template('analysis/analysis.html', 
                           title='Análisis de Requerimiento', 
                           form=form,
                           analisis_obj=analisis_obj,
                           analisis_info=analisis_info,
                           ai_result_data=ai_result_data,
                           ai_result_raw=ai_result_raw_text, 
                           texto_requerimiento=texto_requerimiento_raw,
                           historial_analisis=historial_analisis
                           )


@bp.route('/analysis/clear', methods=['POST'])
@login_required
def clear_analysis():
    flash("Vista limpiada.", "info")
    return redirect(url_for('analysis.analysis_index'))


@bp.route('/analysis/delete/<int:view_id>', methods=['POST'])
@login_required
def delete_analysis(view_id):
    """Elimina un registro de análisis del historial."""
    
    analisis = Analisis.query.get_or_404(view_id)
    if analisis.autor != current_user:
        flash("Acceso no autorizado.", "danger")
        return redirect(url_for('analysis.analysis_index'))
    
    try:
        db.session.delete(analisis)
        db.session.commit()
        flash(f"Análisis '{analisis.nombre_requerimiento}' eliminado.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error al eliminar el análisis: {str(e)}", "danger")
    
    return redirect(url_for('analysis.analysis_index'))


# --- ¡NUEVA RUTA PARA RE-ANALIZAR! ---
@bp.route('/analysis/re-analyze/<int:view_id>', methods=['POST'])
@login_required
def re_analyze(view_id):
    """
    Toma el texto editado del modal, lo re-analiza,
    y actualiza el registro en la base de datos.
    """
    
    # 1. Cargar el análisis existente
    analisis_obj = Analisis.query.get_or_404(view_id)
    if analisis_obj.autor != current_user:
        flash("Acceso no autorizado.", "danger")
        return redirect(url_for('analysis.analysis_index'))
    
    # 2. Obtener el texto editado del formulario del modal
    nuevo_texto = request.form.get('texto_requerimiento')
    
    if not nuevo_texto:
        flash("El texto del requerimiento no puede estar vacío.", "danger")
        return redirect(url_for('analysis.analysis_index', view_id=view_id))
    
    print("\n" + "="*80)
    print(f"🧠 RE-ANALIZANDO REQUERIMIENTO (ID: {view_id})...")
    print("="*80)
    
    # 3. Volver a ejecutar el motor de complejidad
    analisis = analizar_complejidad_requerimiento(nuevo_texto)
    num_casos = analisis['cantidad_casos']
    
    try:
        # 4. Volver a llamar a la API de Gemini (con el modelo y prompt correctos)
        api_key = os.environ.get('GOOGLE_API_KEY') 
        if not api_key:
            flash("Error: GOOGLE_API_KEY no está configurada.", "danger")
            return redirect(url_for('analysis.analysis_index'))
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash-exp') # El modelo que te funciona
        
        # El mismo prompt simplificado
        prompt = f"""
Eres un asistente de QA que genera casos de prueba en formato JSON.
Basado en el siguiente requerimiento, genera EXACTAMENTE {num_casos} casos de prueba.

REQUERIMIENTO A ANALIZAR:
---
{nuevo_texto}
---

INSTRUCCIONES CRÍTICAS:
1. La respuesta debe ser ÚNICAMENTE un array JSON válido, sin texto adicional.
2. Debes generar EXACTAMENTE {num_casos} casos de prueba.
3. Basa tus casos en los Criterios de Aceptación y el texto del requerimiento.
"""
        
        # (Obtener etiquetas de la plantilla)
        plantilla_obj = analisis_obj.plantilla_usada
        mapas = plantilla_obj.mapas.all()
        etiquetas_simples = [m.etiqueta for m in mapas if m.tipo_mapa == 'celda_simple']
        etiquetas_tabla = [m.etiqueta for m in mapas if m.tipo_mapa == 'fila_tabla']

        if etiquetas_simples and etiquetas_tabla:
            prompt += f"""
4. Cada caso debe tener estos campos principales:
   {json.dumps(etiquetas_simples, ensure_ascii=False)}
5. ADICIONALMENTE, cada caso debe tener "PASOS" (array de objetos):
   {json.dumps(etiquetas_tabla, ensure_ascii=False)}
6. Campos principales = strings. "PASOS" = array con 2-10 pasos.
"""
        elif not etiquetas_simples and etiquetas_tabla:
            prompt += f"""
4. Cada caso debe tener estos campos:
   {json.dumps(etiquetas_tabla, ensure_ascii=False)}
5. Todos strings. Si hay campo "Pasos", usar saltos de línea numerados.
"""
        elif etiquetas_simples and not etiquetas_tabla:
            prompt += f"""
4. Cada caso debe tener estos campos:
   {json.dumps(etiquetas_simples, ensure_ascii=False)}
5. Todos strings. Si hay campo "Pasos", usar saltos de línea numerados.
"""

        prompt += f"""
IMPORTANTE: 
- Genera TODOS los campos exactamente como están escritos
- NO inventes campos adicionales
- Asegura que el JSON sea válido y parseable
- EXACTAMENTE {num_casos} casos, ni más ni menos
Responde ÚNICAMENTE con el array JSON: [ ... ]
"""
        
        print("🤖 Enviando (nuevo) prompt a Gemini...")
        response = model.generate_content(prompt)
        
        # 5. Parsear la nueva respuesta
        ai_json_string = None
        casos_generados_count = 0
        try:
            match = re.search(r'\[.*\]', response.text, re.DOTALL)
            if match:
                ai_json_string = match.group(0)
                casos_generados = json.loads(ai_json_string)
                casos_generados_count = len(casos_generados)
                print(f"✅ Re-generación exitosa: {casos_generados_count} casos")
                flash(f"✅ Análisis actualizado: {casos_generados_count} casos generados", "success")
            else:
                flash("⚠️ Respuesta no parseable. Intenta nuevamente.", "warning")
                ai_json_string = response.text
        except json.JSONDecodeError as je:
            print(f"❌ Error JSON: {je}")
            flash("Error: Respuesta de IA inválida.", "danger")
            ai_json_string = response.text

        # 6. ¡ACTUALIZAR (UPDATE) el registro en la BD!
        analisis_obj.texto_requerimiento_raw = nuevo_texto
        analisis_obj.nivel_complejidad = analisis['nivel_complejidad']
        analisis_obj.casos_generados = casos_generados_count
        analisis_obj.criterios_detectados = analisis['metricas']['criterios_aceptacion']
        analisis_obj.palabras_analizadas = analisis['metricas']['palabras']
        analisis_obj.ai_result_json = ai_json_string
        
        db.session.commit() # Guardar los cambios en el objeto existente
        
        print(f"💾 Análisis {analisis_obj.id} actualizado en la base de datos.")
        
    except Exception as e:
        flash(f"Error con API Gemini: {str(e)}", "danger")
        traceback.print_exc()
    
    # 7. Redirigir de vuelta a la vista del análisis
    return redirect(url_for('analysis.analysis_index', view_id=view_id))


@bp.route('/generate_file/<int:view_id>')
@login_required
def generate_file(view_id):
    """
    Genera y descarga el archivo entregable para un análisis específico.
    (Sin cambios)
    """
    
    analisis_obj = Analisis.query.get_or_404(view_id)
    if analisis_obj.autor != current_user:
        flash("Acceso no autorizado.", "danger")
        return redirect(url_for('analysis.analysis_index'))

    json_text = analisis_obj.ai_result_json
    plantilla_obj = analisis_obj.plantilla_usada
    
    if not json_text or not plantilla_obj:
        flash("Error: Faltan datos de JSON o plantilla en este análisis.", "danger")
        return redirect(url_for('analysis.analysis_index', view_id=view_id))
        
    try:
        match = re.search(r'\[.*\]', json_text, re.DOTALL)
        datos_ia = json.loads(match.group(0))
        print(f"📦 Generando archivo para Análisis {view_id} con {len(datos_ia)} casos")
    except Exception as e:
        flash("Error: Datos JSON corruptos en este análisis.", "danger")
        traceback.print_exc()
        return redirect(url_for('analysis.analysis_index', view_id=view_id))
        
    buffer_memoria = None
    mimetype = ""
    download_name = ""
    
    if plantilla_obj.tipo_archivo == 'Excel':
        buffer_memoria = generar_excel_entregable(plantilla_obj, datos_ia)
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        download_name = f"entregable_{analisis_obj.id}_{plantilla_obj.nombre_plantilla}.xlsx"
        
    elif plantilla_obj.tipo_archivo == 'Word':
        buffer_memoria = generar_word_entregable(plantilla_obj, datos_ia)
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        download_name = f"entregable_{analisis_obj.id}_{plantilla_obj.nombre_plantilla}.docx"
    
    if buffer_memoria:
        return send_file(
            buffer_memoria,
            as_attachment=True,
            download_name=download_name,
            mimetype=mimetype
        )
    else:
        flash(f"Error generando {plantilla_obj.tipo_archivo}. Revisa logs.", "danger")
        return redirect(url_for('analysis.analysis_index', view_id=view_id))